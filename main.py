import ast
import streamlit as st
from ollama import Client
import io

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("The 'python-docx' package is required for .docx export. Please install it with 'pip install python-docx'")
    
# Add these imports at the top
import os
import mimetypes
from pathlib import Path

def detect_file_type(filename):
    """Detect file type based on extension and content."""
    if not filename:
        return "unknown"
    
    ext = Path(filename).suffix.lower()[1:]  # Get extension without dot
    mime_type, _ = mimetypes.guess_type(filename)
    
    # Code files
    code_exts = ['py', 'js', 'jsx', 'ts', 'tsx', 'java', 'c', 'cpp', 'h', 'hpp', 
                'cs', 'go', 'rs', 'rb', 'php', 'sh', 'pl', 'r', 'm', 'jl']
    if ext in code_exts:
        return 'code'
        
    # Data files
    data_exts = ['json', 'yaml', 'yml', 'xml', 'csv', 'toml', 'ini', 'cfg']
    if ext in data_exts:
        return 'data'
        
    # Document files
    doc_exts = ['md', 'txt', 'html', 'htm', 'css', 'pdf', 'doc', 'docx']
    if ext in doc_exts or (mime_type and any(t in mime_type for t in ['text', 'document'])): 
        return 'document'
        
    return 'unknown'

def read_file_content(file):
    """Read file content based on its type."""
    try:
        if file.type and ('text/' in file.type or file.type in ['application/json', 'application/xml']):
            return file.getvalue().decode('utf-8')
        else:
            # For binary files, we'll just note the file type
            return f"[Binary file: {file.name}, Type: {file.type or 'unknown'}]"
    except Exception as e:
        return f"Error reading file: {str(e)}"

# Helper
def extract_elements(code):
    """Extracts all top-level classes and functions with docstrings and source code."""
    tree = ast.parse(code)
    elements = []
    for node in tree.body:
        if isinstance(node, (ast.FunctionDef, ast.ClassDef, ast.AsyncFunctionDef)):
            # Get basic info
            element_type = node.__class__.__name__
            name = node.name
            doc = ast.get_docstring(node) or ""
            
            # Get full source code
            src = ast.get_source_segment(code, node) or ""
            
            # Get line numbers
            start_line = getattr(node, 'lineno', 0)
            end_line = getattr(node, 'end_lineno', start_line)
            
            # Get function arguments if it's a function
            args = []
            if hasattr(node, 'args'):
                args = [arg.arg for arg in node.args.args]
            
            elements.append({
                'type': element_type.replace('Def', ''),  # 'FunctionDef' -> 'Function', 'ClassDef' -> 'Class'
                'name': name,
                'docstring': doc,
                'source': src,
                'start_line': start_line,
                'end_line': end_line,
                'args': args if args else None,
                'has_return': any(isinstance(n, ast.Return) for n in ast.walk(node))
            })
    return elements

def llama_explain(code_elements):
    """
    Send code info to Llama via Ollama and get detailed non-technical summary.
    
    The prompt is structured as follows:
    
    [STATIC] - These parts never change:
    - The initial instructions to the AI about its role and how to respond
    - The section headers (e.g., 'Documentation:', 'Code:')
    - The formatting of the response
    
    [DYNAMIC] - These parts are filled with actual code analysis:
    - Function/Class names and types
    - Line numbers where code appears
    - Documentation strings from the code
    - Function arguments (if any)
    - Return value indicators
    - The actual source code
    
    Example of how the prompt will look:
    
    [SYSTEM PROMPT - Static]
    You are a helpful assistant. Read the information below and summarize what this Python code does,
    explaining it in detail, in plain, non-technical English, to a non-programmer. Avoid jargon.
    Where possible, use analogies and concrete examples.
    
    Here is the code to summarize:
    
    [DYNAMIC CONTENT - Example for one function]
    Function 'calculate_total':
    Location: Lines 5-10
    Documentation: Calculates the total price including tax
    Arguments: price, tax_rate
    Returns: Yes
    Code:
    def calculate_total(price, tax_rate):
    """
    # Compose prompt with enhanced information
    combined = "\n\n".join(
        f"{el['type']} '{el['name']}':\n"
        f"Location: Lines {el['start_line']}-{el['end_line']}\n"
        f"Documentation: {el['docstring']}\n"
        + (f"Arguments: {', '.join(el['args'])}\n" if el['args'] else "") 
        + ("Returns: Yes\n" if el['has_return'] and el['type'] != 'Class' else "")
        + f"Code:\n{el['source'] or ''}"
        for el in code_elements
    )
    
    system_prompt = (
    "You are a helpful assistant. Your job is to explain the Python code and workflow described below "
    "in plain, non-technical English to someone without a programming background. Avoid technical jargon. "
    "Use relatable analogies and simple examples where appropriate.\n\n"

    "📌 Task Overview:\n"
    "- Break down the Python code into understandable parts.\n"
    "- Present the explanation in a two-column table:\n"
    "    1. Section of the prompt\n"
    "    2. Whether it is dynamic or static\n\n"

    "🧠 User Context:\n"
    "- The user is building an automated assessment tool using LLMs.\n"
    "- The tool generates subtopics and test questions from a topic + grade + learning objective.\n"
    "- They want help modifying Prompt 1 to include a new variable: the learning objective.\n\n"

    "💡 Input Example:\n"
    "    topic = 'Ratios and Proportional Relationships'\n"
    "    student_class = '6th standard'\n"
    "    learning_objective = 'Understand ratio concepts and use ratio reasoning to solve problems.'\n\n"

    "📝 Full Prompt 1 (Subtopic Generator):\n"
    "I want a list of sub-topics for the topic \"{topic}\" which is taught to a \"{student_class}\" student "
    "with a learning objective \"{learning_objective}\".\n"
    "First, output the learning objective exactly as given.\n"
    "Then, output the sub-topics ONLY as a Python list. Do not include any commentary or explanation.\n\n"

    "🔧 System Context:\n"
    "We are building an automated assessment web app where questions are usually uploaded manually into a MySQL database. "
    "This tool uses large language models to generate those questions automatically, saving time and effort.\n\n"

    "📋 Additional User Requests:\n"
    "- Modify the original code to include the new learning objective variable.\n"
    "- Ensure that generate_questions_for_subtopic also uses the learning objective.\n"
    "- Recreate the dynamic/static breakdown table of Prompt 1 and Prompt 2.\n"
    "- Show an example of what Prompt 1 and Prompt 2 look like after the code runs.\n"
    "- Give a step-by-step guide to feeding these prompts into a ChatGPT conversation.\n"
    "- Convert the instructions into clean documentation.\n"
    "- Provide a downloadable .docx version of the documentation.\n\n"

    f"{combined}\n\n"
    "🧾 Now, please provide a detailed, beginner-friendly explanation:"
)


    try:
        client = Client(host='http://localhost:11434')
        response = client.chat(model='llama2', messages=[{"role": "user", "content": system_prompt}])
    
    except Exception as e:
        return f"Error generating explanation: {str(e)}"
    return response['message']['content']

# ---- Streamlit UI ----
# ---- Streamlit UI ----
st.title("File Analyzer & Summarizer")

uploaded_file = st.file_uploader(
    "Upload any file for analysis", 
    type=None,  # Accept all file types
    accept_multiple_files=False,
    help="Upload any text-based file (code, documents, data files) for analysis"
)

if uploaded_file:
    try:
        file_type = detect_file_type(uploaded_file.name)
        file_content = read_file_content(uploaded_file)
        
        st.header(f"File Analysis: {uploaded_file.name}")
        st.write(f"**Type:** {file_type.capitalize()} file")
        
        if file_type == 'code':
            # Existing code analysis for Python files
            if uploaded_file.name.endswith('.py'):
                code_elements = extract_elements(file_content)
                # ... rest of the code analysis logic ...
            else:
                st.subheader("File Content")
                st.code(file_content, language='text')
                
        elif file_type == 'data':
            st.subheader("Data Content")
            st.json(file_content) if uploaded_file.name.endswith('.json') else st.code(file_content)
            
        elif file_type == 'document':
            st.subheader("Document Content")
            st.text_area("Content", file_content, height=300)
            
        else:
            st.warning("This file type is not fully supported for analysis.")
            st.download_button(
                label="Download File",
                data=uploaded_file,
                file_name=uploaded_file.name,
                mime=uploaded_file.type
            )
            
        # Add a button to generate DOCX report
        if st.button("Generate DOCX Report"):
            doc = Document()
            doc.add_heading(f"File Analysis Report: {uploaded_file.name}", 0)
            doc.add_paragraph(f"File Type: {file_type.capitalize()}")
            
            if file_type == 'code' and uploaded_file.name.endswith('.py'):
                doc.add_heading("Code Structure:", level=1)
                for el in code_elements:
                    doc.add_heading(f"{el['type']}: {el['name']} (Lines {el['start_line']}-{el['end_line']})", level=2)
                    if el['args']:
                        doc.add_paragraph(f"Arguments: {', '.join(el['args'])}")
                    if el['type'] != 'Class' and el['has_return']:
                        doc.add_paragraph("Returns: Yes")
                    if el['docstring']:
                        doc.add_paragraph(f"Documentation:\\n{el['docstring']}")
                    doc.add_paragraph("Source Code:")
                    doc.add_paragraph(el['source'], style='Intense Quote')
            else:
                doc.add_heading("File Content:", level=1)
                doc.add_paragraph(file_content)
            
            # Save to a BytesIO buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Create download button
            st.download_button(
                label="Download DOCX Report",
                data=buffer,
                file_name=f"analysis_{Path(uploaded_file.name).stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
else:
    st.info("Please upload a file to analyze")