"""
Document generation module for creating Word documents from code analysis.
"""
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO
from dataclasses import dataclass
import json
from ollama import Client

@dataclass
class CodeElementSummary:
    """Class to hold simplified information about a code element."""
    name: str
    type: str
    description: str
    example: str = ""

def generate_non_tech_explanation(code_elements: List[Dict[str, Any]], 
                               model: str = 'llama2',
                               host: str = 'http://localhost:11434') -> Tuple[str, List[CodeElementSummary]]:
    """
    Generate a non-technical explanation of the code using LLM.
    
    Args:
        code_elements: List of code element dictionaries from the analyzer
        model: The Ollama model to use for generation
        host: The Ollama server host
        
    Returns:
        Tuple of (general_explanation, element_descriptions)
    """
    try:
        client = Client(host=host)
        
        # Create a simplified version of code elements for the prompt
        simplified_elements = []
        for el in code_elements:
            simplified = {
                'name': el.get('name', 'unnamed'),
                'type': el.get('type', 'code'),
                'docstring': el.get('docstring', ''),
                'args': el.get('args', []),
                'has_return': el.get('has_return', False),
                'source': el.get('source', '')[:500]  # Limit source length
            }
            simplified_elements.append(simplified)
        
        # Prepare the prompt
        system_prompt = """You are a helpful assistant that explains code in simple, non-technical terms. 
        Your audience has little to no programming knowledge. Use analogies and simple language.
        Explain what the code does, not how it does it."""
        
        user_prompt = f"""Please explain the following Python code elements in a way that's easy for non-programmers to understand.
        For each element, provide:
        1. A simple explanation of what it does
        2. A real-world analogy
        3. A simple example of how it might be used
        
        Code elements to explain: {json.dumps(simplified_elements, indent=2)}
        
        Format your response as a JSON object with these fields:
        - general_overview: A brief overview of what the code does as a whole
        - elements: A list of objects, each with 'name', 'type', 'explanation', 'analogy', and 'example' fields
        """
        
        # Get the LLM response
        response = client.chat(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            format="json"
        )
        
        # Parse the response
        try:
            result = json.loads(response['message']['content'])
            
            # Format the general explanation
            general = f"# Understanding the Code: A Friendly Guide\n\n"
            general += f"## What This Code Does\n\n{result.get('general_overview', 'This code performs specific operations.')}\n\n"
            
            # Format the elements
            elements = []
            for elem in result.get('elements', []):
                desc = f"{elem.get('explanation', '')}\n\n"
                if 'analogy' in elem:
                    desc += f"*Think of it like:* {elem['analogy']}\n\n"
                elements.append(CodeElementSummary(
                    name=elem.get('name', 'unnamed'),
                    type=elem.get('type', 'code').capitalize(),
                    description=desc.strip(),
                    example=elem.get('example', '')
                ))
                
                # Add to general explanation
                general += f"### {elem.get('type', 'Element').capitalize()}: {elem.get('name', 'Unnamed')}\n"
                general += f"{desc}\n"
                if 'example' in elem:
                    general += f"*Example:* {elem['example']}\n\n"
            
            # Add conclusion
            general += """
## Why This Is Helpful

1. **For Beginners**: Understand code without needing to learn programming first
2. **For Teams**: Helps non-technical team members understand technical work
3. **For Documentation**: Creates clear, accessible records of what code does

## Real-World Impact

This kind of code documentation helps bridge the gap between technical and non-technical stakeholders, making technology more accessible to everyone.
"""
            
            return general, elements
            
        except json.JSONDecodeError:
            # Fallback to simple explanation if JSON parsing fails
            return _generate_fallback_explanation(code_elements)
            
    except Exception as e:
        print(f"Error generating LLM explanation: {e}")
        return _generate_fallback_explanation(code_elements)

def _generate_fallback_explanation(code_elements: List[Dict[str, Any]]) -> Tuple[str, List[CodeElementSummary]]:
    """Generate a simple fallback explanation without LLM."""
    elements = []
    general = "# Code Explanation\n\nThis code contains the following components:\n\n"
    
    for el in code_elements:
        el_type = el.get('type', 'code')
        name = el.get('name', 'unnamed')
        desc = f"{el_type.capitalize()} '{name}'"
        if el.get('docstring'):
            desc += f": {el['docstring']}"
            
        elements.append(CodeElementSummary(
            name=name,
            type=el_type,
            description=desc,
            example=""
        ))
        
        general += f"- {desc}\n"
    
    general += "\nFor a more detailed explanation, please check the technical documentation."
    return general, elements

def generate_codebase_summary(code_elements: List[Dict[str, Any]], 
                           model: str = 'llama2',
                           host: str = 'http://localhost:11434') -> Dict[str, Any]:
    """
    Generate a comprehensive summary of the codebase using LLM.
    
    Args:
        code_elements: List of code element dictionaries from the analyzer
        model: The Ollama model to use for generation
        host: The Ollama server host
        
    Returns:
        Dictionary containing the summary with the following keys:
        - overview: General overview of the codebase
        - architecture: High-level architecture description
        - key_components: List of main components
        - data_flow: Description of how data moves through the system
        - dependencies: External libraries and services used
        - setup_instructions: How to set up the project
        - usage_examples: Example usage of the main components
    """
    try:
        client = Client(host=host)
        
        # Prepare the prompt
        system_prompt = """You are a senior software architect analyzing a codebase. 
        Provide a comprehensive summary in JSON format with the following structure:
        {
            "overview": "Brief overview of what the codebase does",
            "architecture": "High-level architecture description",
            "key_components": ["List", "of", "main", "components"],
            "data_flow": "Description of how data moves through the system",
            "dependencies": ["List", "of", "main", "dependencies"],
            "setup_instructions": "Step-by-step setup instructions",
            "usage_examples": "Example usage of the main components"
        }"""
        
        # Create a simplified version of code elements for the prompt
        simplified_elements = []
        for el in code_elements:
            simplified = {
                'name': el.get('name', 'unnamed'),
                'type': el.get('type', 'code'),
                'docstring': el.get('docstring', ''),
                'file': el.get('file', 'unknown')
            }
            simplified_elements.append(simplified)
        
        user_prompt = f"""Analyze the following code elements and provide a comprehensive summary:
        {json.dumps(simplified_elements, indent=2)}
        
        Focus on:
        1. The overall purpose of the codebase
        2. How the different components interact
        3. The main data flows
        4. Key dependencies
        5. How to set up and use the system
        """
        
        # Get the LLM response
        response = client.chat(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            format="json"
        )
        
        # Parse and return the response
        try:
            return json.loads(response['message']['content'])
        except json.JSONDecodeError:
            return {
                "error": "Failed to parse the LLM response",
                "raw_response": response['message']['content']
            }
            
    except Exception as e:
        return {
            "error": f"Error generating codebase summary: {str(e)}"
        }

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def create_document(code_elements: List[Dict[str, Any]], explanation: str, 
                   title: str = "Python Code Analysis",
                   include_non_tech: bool = True,
                   include_summary: bool = True,
                   model: str = 'llama2',
                   host: str = 'http://localhost:11434') -> Optional[BytesIO]:
    """
    Create a Word document from code analysis and explanation.
    
    Args:
        code_elements: List of code element dictionaries, each can include 'file' key
        explanation: Generated explanation text
        title: Title for the document
        include_non_tech: Whether to include non-technical explanation
        include_summary: Whether to include a comprehensive codebase summary
        model: The Ollama model to use for non-technical explanations
        host: The Ollama server host
        
    Returns:
        BytesIO buffer containing the document, or None if docx is not available
    """
    if not DOCX_AVAILABLE:
        return None
        
    # Generate non-technical explanation if requested
    non_tech_explanation = ""
    if include_non_tech and code_elements:
        non_tech_explanation, _ = generate_non_tech_explanation(
            code_elements, 
            model=model,
            host=host
        )
        
    # Generate codebase summary if requested
    codebase_summary = {}
    if include_summary and code_elements:
        codebase_summary = generate_codebase_summary(
            code_elements,
            model=model,
            host=host
        )
        
    def _get_or_create_code_style(doc):
        """Get the Code style or create it if it doesn't exist."""
        try:
            return doc.styles['Code']
        except KeyError:
            code_style = doc.styles.add_style('Code', 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(10)
            code_style.paragraph_format.space_after = Pt(6)
            return code_style
        
    try:
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(level=0)
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Ensure we have the Code style
        _get_or_create_code_style(doc)
        
        # Add non-technical explanation if available
        if non_tech_explanation:
            doc.add_heading("Non-Technical Overview", level=1)
            lines = non_tech_explanation.split('\n')
            current_paragraph = None
            
            for line in lines:
                if line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:].strip(), level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:].strip(), level=4)
                elif line.strip() == '':
                    if current_paragraph:
                        current_paragraph.add_run('\n')
                else:
                    current_paragraph = doc.add_paragraph(line)
                
            doc.add_page_break()
            
        # Add technical summary section
        doc.add_heading("Technical Summary", level=1)
        
        # Count elements by type and file
        elements_by_file = {}
        elements_by_type = {}
        
        for el in code_elements:
            file_name = el.get('file', 'unknown.py')
            el_type = el.get('type', 'Unknown')
            
            # Count by file
            if file_name not in elements_by_file:
                elements_by_file[file_name] = {'functions': 0, 'classes': 0, 'async_functions': 0}
            
            if el_type == 'Function':
                elements_by_file[file_name]['functions'] += 1
            elif el_type == 'AsyncFunction':
                elements_by_file[file_name]['async_functions'] += 1
            elif el_type == 'Class':
                elements_by_file[file_name]['classes'] += 1
            
            # Count by type
            if el_type not in elements_by_type:
                elements_by_type[el_type] = 0
            elements_by_type[el_type] += 1
        
        # Add summary table
        if elements_by_file:
            doc.add_heading("Files Analyzed", level=2)
            for file, counts in elements_by_file.items():
                parts = []
                if counts['classes']:
                    parts.append(f"{counts['classes']} classes")
                if counts['functions']:
                    parts.append(f"{counts['functions']} functions")
                if counts['async_functions']:
                    parts.append(f"{counts['async_functions']} async functions")
                doc.add_paragraph(f"• {file}: {', '.join(parts) if parts else 'No elements found'}")
        
        # Add element type summary
        if elements_by_type:
            doc.add_paragraph(f"Total elements found: {sum(elements_by_type.values())}")
            for el_type, count in elements_by_type.items():
                doc.add_paragraph(f"• {el_type}s: {count}", style='List Bullet')
        
        # Add code structure section
        doc.add_heading("Code Structure", level=1)
        
        # Group elements by file
        elements_by_file = {}
        for el in code_elements:
            file_name = el.get('file', 'unknown.py')
            if file_name not in elements_by_file:
                elements_by_file[file_name] = []
            elements_by_file[file_name].append(el)
        
        # Process each file
        for file_name, elements in elements_by_file.items():
            # Add file header
            doc.add_heading(f"File: {file_name}", level=2)
            
            for el in elements:
                # Add element header
                el_header = f"{el['type']}: {el['name']} (Lines {el['start_line']}-{el['end_line']})"
                doc.add_heading(el_header, level=3)
                
                # Add metadata
                if el['args']:
                    doc.add_paragraph(f"Arguments: {', '.join(el['args'])}")
                    
                if el['type'] != 'Class' and el['has_return']:
                    doc.add_paragraph("Returns: Yes")
                    
                if el['docstring']:
                    doc.add_paragraph("Documentation:")
                    doc.add_paragraph(el['docstring'], style='Intense Quote')
                
                # Add source code
                doc.add_paragraph("Source Code:")
                doc.add_paragraph(el['source'], style='Code')
                
                # Add a small separator between elements
                doc.add_paragraph()
                doc.add_paragraph("-" * 40)
                doc.add_paragraph()
        
        # Add explanation section
        doc.add_heading("AI Explanation", level=1)
        doc.add_paragraph(explanation)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except Exception as e:
        raise RuntimeError(f"Error generating document: {str(e)}")
